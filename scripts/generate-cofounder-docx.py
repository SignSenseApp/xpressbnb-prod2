"""Generate cofounder update DOCX (Hinglish)."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).resolve().parents[1] / "XpressBnB_Cofounder_Update_May2026.docx"

NAVY = RGBColor(0x0A, 0x1F, 0x44)
GOLD = RGBColor(0xC9, 0xA2, 0x27)
MUTED = RGBColor(0x55, 0x55, 0x55)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = NAVY


def add_body(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = MUTED if not bold else NAVY
    run.bold = bold
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.25


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = MUTED
    p.paragraph_format.space_after = Pt(4)


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = title.add_run("XpressBnB — Cofounder Update")
    t_run.bold = True
    t_run.font.size = Pt(26)
    t_run.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s_run = sub.add_run("May 2026  •  Full-stack progress report")
    s_run.font.size = Pt(12)
    s_run.font.color.rgb = GOLD
    s_run.italic = True

    doc.add_paragraph()

    add_body(
        doc,
        "Bhai / Partner —",
        bold=True,
    )
    add_body(
        doc,
        "Maine ye note isliye likha hai taaki tum clearly dekh sako: XpressBnB ab \"idea on paper\" nahi raha — "
        "ye live product hai, daily improve ho raha hai, aur hum dono ke liye ek solid base ready hai. "
        "Neeche sab kuch simple Hinglish mein hai — technical jargon kam, impact zyada.",
    )

    add_heading(doc, "1. Big Picture — Hum Kahan Khade Hain", 2)
    add_body(
        doc,
        "XpressBnB ab production par chal raha hai. Matlab real users, real bookings flow, real brand — "
        "sirf demo nahi. Maine last phase mein product ko polish kiya: listings sahi dikhni chahiye, "
        "login ke baad sahi homepage, mobile par premium feel, Google par sahi branding, aur honest reviews — "
        "taaki trust build ho.",
    )
    add_body(
        doc,
        "Short mein: platform ab investor / host / guest ko dikhane layak hai. "
        "Ab humara focus growth + operations par shift ho sakta hai — tumhari expertise yahan game-changer hogi.",
        bold=True,
    )

    add_heading(doc, "2. Live Links (Abhi Check Kar Sakte Ho)", 2)
    add_bullet(doc, "Production site: https://xpressbnb.vercel.app")
    add_bullet(doc, "Cofounder / demo codebase (tumhara dedicated repo): https://github.com/SignSenseApp/xpressbnb-prod2")
    add_bullet(doc, "Main company repo: https://github.com/SignSenseApp/xpressbnb.com")

    add_heading(doc, "3. Maine Kya-Kya Deliver Kiya (Last Sprint)", 2)

    add_heading(doc, "A. Product & UX — Guest Side", 3)
    add_bullet(doc, "Naya homepage design live — clean, Apple-style spacing, mobile-first.")
    add_bullet(doc, "Gurgaon / Gurugram / Delhi / Rishikesh listings fix — pehle properties miss ho rahi thi; ab city matching smart hai.")
    add_bullet(doc, "Login / logout ke baad galat purana map UI nahi aata — user hamesha naye homepage par land hota hai.")
    add_bullet(doc, "Fake \"113 reviews\" hata diye — ab sirf real review count dikhta hai (trust ke liye bahut important).")

    add_heading(doc, "B. Brand & Trust", 3)
    add_bullet(doc, "Favicon + app icons + manifest — Google search aur mobile home screen par proper XpressBnB logo.")
    add_bullet(doc, "SEO metadata update — brand consistently XpressBnB dikhe.")

    add_heading(doc, "C. Engineering & Deployment", 3)
    add_bullet(doc, "Vercel production deploy — site live, updates push karne par auto-build.")
    add_bullet(doc, "Alag GitHub repo (xpressbnb-prod2) tumhare liye setup — tum bina main codebase touch kiye apna version / demo chala sakte ho.")
    add_bullet(doc, "Saari PNG / image assets repo mein push — koi file missing nahi.")
    add_bullet(doc, "Payments (Razorpay), host dashboard, calendar, maps, Supabase backend — pehle se integrated; ab polish layer add hui hai.")

    add_heading(doc, "4. Tumhare Liye Kya Ready Hai", 2)
    add_body(
        doc,
        "Maine intentionally tumhare liye alag repo banaya hai taaki:",
    )
    add_bullet(doc, "Tum freely explore kar sako, deploy kar sako, investors ko dikha sako.")
    add_bullet(doc, "Live production (xpressbnb.vercel.app) disturb na ho.")
    add_bullet(doc, "Hum parallel kaam kar saken — main product/engineering, tum growth / partnerships / ops.")

    add_heading(doc, "5. Mujhe Tumse Kya Chahiye (Clear Ask)", 2)
    add_body(doc, "Ye sprint tabhi 10x banega jab tum apna superpower lagaoge. Please ye 4 cheezein priority par:", bold=True)
    add_bullet(doc, "Repo clone karo: git clone https://github.com/SignSenseApp/xpressbnb-prod2.git — locally chalao (npm install → npm run dev).")
    add_bullet(doc, "5 hosts / property owners se baat karo — Gurgaon + Rishikesh se start; unhe live link bhejo.")
    add_bullet(doc, "1-pager pitch ready karo (screenshots + live URL) — investors / partners ke liye.")
    add_bullet(doc, "Weekly 30-min sync fix karo — blockers turant solve karenge.")

    add_heading(doc, "6. Agla Phase — Hum Dono Milke", 2)
    add_bullet(doc, "Cofounder repo ko Vercel par alag preview URL (optional) — tumhara personal demo link.")
    add_bullet(doc, "Google Search Console — favicon indexing (brand visibility).")
    add_bullet(doc, "Host onboarding funnel + first 10 paid listings target.")
    add_bullet(doc, "Marketing / Instagram / WhatsApp campaigns — yahan tum lead karoge, main product support dunga.")

    add_heading(doc, "7. Closing — Seedha Baat", 2)
    add_body(
        doc,
        "Main product side se full commitment de raha hoon — late nights, bug fixes, deploys, details. "
        "Tum dekh sakte ho: codebase pushed hai, site live hai, issues fix ho chuke hain. "
        "Ab ball tumhare court mein hai — outreach, partnerships, hustle. "
        "Hum dono alag skills le kar aaye hain; ab execution ka time hai.",
    )
    add_body(
        doc,
        "Agar kuch samajh na aaye ya deploy karna ho — message karo, 15 min call par sab set kar denge. "
        "Chalo isko NCR + Uttarakhand ka strongest boutique stays brand banate hain. 🚀",
        bold=True,
    )

    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig_run = sig.add_run("— Your cofounder & builder\nXpressBnB Team")
    sig_run.font.size = Pt(11)
    sig_run.font.color.rgb = NAVY
    sig_run.italic = True

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f_run = footer.add_run("Confidential — for cofounder eyes only")
    f_run.font.size = Pt(9)
    f_run.font.color.rgb = MUTED
    f_run.italic = True

    doc.save(OUT)
    print(f"Created: {OUT}")


if __name__ == "__main__":
    build()
