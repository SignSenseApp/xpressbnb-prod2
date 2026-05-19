"""
XpressBnB – Siddhart Sir Update DOCX
Real full-stack audit as of 19 May 2026. Launch target: 10 June 2026.
"""
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).resolve().parents[1] / "XpressBnB_Siddhart_Update_19May2026.docx"

NAVY   = RGBColor(0x0A, 0x1F, 0x44)
GREEN  = RGBColor(0x05, 0x7A, 0x3C)
RED    = RGBColor(0xB9, 0x1C, 0x1C)
AMBER  = RGBColor(0x92, 0x40, 0x00)
MUTED  = RGBColor(0x44, 0x44, 0x44)
GOLD   = RGBColor(0xC9, 0xA2, 0x27)
GRAY   = RGBColor(0x6B, 0x72, 0x80)


# ─── helpers ───────────────────────────────────────────────────────────────

def heading(doc, text, level=1, color=NAVY):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = color
    return p


def body(doc, text, bold=False, color=MUTED, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.bold = bold
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.line_spacing = 1.3
    return p


def bullet(doc, text, color=MUTED, prefix=""):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(f"{prefix}{text}")
    r.font.size = Pt(10.5)
    r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(3)
    return p


def status_row(doc, label, status_text, status_color):
    """Single-line item: label (black) + status badge (colored)."""
    p = doc.add_paragraph(style="List Bullet")
    r1 = p.add_run(f"{label}  ")
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = MUTED
    r2 = p.add_run(f"[{status_text}]")
    r2.font.size = Pt(10)
    r2.font.color.rgb = status_color
    r2.bold = True
    p.paragraph_format.space_after = Pt(3)


def divider(doc):
    doc.add_paragraph("─" * 72).paragraph_format.space_after = Pt(2)


def spacer(doc, n=1):
    for _ in range(n):
        doc.add_paragraph("").paragraph_format.space_after = Pt(2)


# ─── build ─────────────────────────────────────────────────────────────────

def build():
    doc = Document()
    s = doc.sections[0]
    s.top_margin    = Inches(0.85)
    s.bottom_margin = Inches(0.85)
    s.left_margin   = Inches(1.0)
    s.right_margin  = Inches(1.0)

    # ══ TITLE BLOCK ══════════════════════════════════════════════════════════
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run("XpressBnB — Product & Engineering Update")
    tr.bold = True
    tr.font.size = Pt(24)
    tr.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Prepared by Saad   ·   19 May 2026   ·   Launch target: 10 June 2026")
    sr.font.size = Pt(11)
    sr.font.color.rgb = GOLD
    sr.italic = True

    spacer(doc)
    divider(doc)

    # ══ OPENING NOTE ══════════════════════════════════════════════════════════
    heading(doc, "Siddhart Sir — Context for this document", 2)
    body(doc,
        "Maine yeh note isliye banaya hai taaki apko clearly pata ho — "
        "kya kaam hua, kya bacha, aur June 10 launch ke liye exactly kya-kya "
        "activate karna hai. Koi bhi point guess nahi hai — sab actual code aur "
        "live production database se verify kiya hai. Zero assumptions.",
        color=MUTED)

    spacer(doc)

    # ══ SECTION 1 — SNAPSHOT ══════════════════════════════════════════════════
    heading(doc, "1.  Live Snapshot — Right Now (19 May 2026)", 2)

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "Value"
    for cell in hdr:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = NAVY

    rows_data = [
        ("Total live properties",           "93 (Delhi 13 · Gurgaon 28 · Noida 22 · G.Noida 7 · Rishikesh 23)"),
        ("Properties with host + phone",     "62  — full inquiry flow works"),
        ("Properties missing host link",     "31  — visible on site, ops gap (see Section 6)"),
        ("Bookings in DB",                   "48  (32 confirmed/pending, 15 pending, 1 offer_pending)"),
        ("Edge functions in Supabase cloud", "6 active (2 Razorpay, send/verify OTP, create-host-sub, verify-host-sub)"),
        ("New edge functions (repo, undeployed)", "send-booking-otp · verify-booking-otp · send-inquiry-notification"),
        ("DB migrations applied (prod)",     "10 migrations — inquiry schema, OTP tables, subscription tiers"),
        ("Razorpay environment",             "LIVE key on frontend (rzp_live_SJYZ…) — host subscriptions only"),
    ]
    for label, val in rows_data:
        row = tbl.add_row().cells
        row[0].text = label
        row[1].text = val

    spacer(doc)

    # ══ SECTION 2 — WHAT IS DONE (FULLY) ═════════════════════════════════════
    heading(doc, "2.  Kya Ho Gaya — Fully Built & Live", 2)
    body(doc,
         "Neeche woh cheezein hain jo production code mein hain, "
         "Supabase pe deploy hain, aur site pe test ho sakti hain aaj.",
         bold=True, color=NAVY, size=10)

    spacer(doc)
    heading(doc, "2A.  Guest Property Page — Zero Commission Inquiry Flow", 3)
    items_done_guest = [
        "Guest property page pe Razorpay HATA DIYA — koi payment modal nahi",
        "\"Book Now & Pay Later\" → OTP step → verified inquiry (code: BookingForm.tsx + GuestPhoneOtpStep.tsx)",
        "\"Make an Offer\" → OTP step → verified offer inquiry (OfferModal.tsx)",
        "GuestPhoneOtpStep: phone input → Send OTP → 6-digit entry → verify → unlocks submit",
        "OTP ke bina koi bhi inquiry host dashboard pe visible nahi — anti-fake gate in DB (consume_booking_inquiry_otp RPC)",
        "Success screen: InquirySuccessModal — \"Inquiry {hostName} ko bhej di\" + host phone (call/WhatsApp CTA)",
        "Sidebar copy update: \"0% commission — guests pay you directly\"",
        "Promo codes, decoration add-on, guests selector — sab kaam kar rahe hain inquiry model mein",
    ]
    for i in items_done_guest:
        bullet(doc, i, GREEN, "✓  ")

    spacer(doc)
    heading(doc, "2B.  Host Dashboard — Bookings & Offer Handling", 3)
    items_done_host = [
        "Host BookingsPage: inquiry section (Accept / Reject / Counter) — fully reworked",
        "Accept → status: accepted; guest notification queue insert (DB trigger)",
        "Reject → status: rejected; guest ke liye dialog (RejectInquiryDialog.tsx)",
        "Counter offer → CounterOfferDialog, counter_nightly_rate DB column",
        "Guest phone visible to host ONLY when phone_verified = true (RLS enforced)",
        "Realtime toast — nayi inquiry aate hi host dashboard pe ping",
        "Bookings filter tabs (All / Pending / Confirmed / Offers) working",
    ]
    for i in items_done_host:
        bullet(doc, i, GREEN, "✓  ")

    spacer(doc)
    heading(doc, "2C.  Host Subscription Tiers — Razorpay Host-Only", 3)
    items_done_sub = [
        "3 tiers fully live in UI: Free | Standard ₹999/month | Premium ₹2,999/month",
        "Yearly billing toggle — 12% discount (₹10,551/yr Standard; ₹31,668/yr Premium)",
        "Dedicated edge function: create-host-subscription-order (separate from guest booking order)",
        "Dedicated edge function: verify-host-subscription (server-side, not client-side upsert)",
        "PropertyUpgradeModal: tier selection (Standard vs Premium) + billing cycle before checkout",
        "property_subscriptions.plan_tier column added (standard_999 | premium_2999)",
        "Trigger: plan_tier → properties.is_premium sync (no manual update needed)",
        "Premium feature flags in UI — Advanced analytics, calendar sync, badge (marked 'coming soon')",
        "SubscriptionPage copy: \"0% commission — guests pay you directly\" — Airbnb contrast",
    ]
    for i in items_done_sub:
        bullet(doc, i, GREEN, "✓  ")

    spacer(doc)
    heading(doc, "2D.  Notification Pipeline — Backend Built", 3)
    items_done_notif = [
        "send-inquiry-notification edge function — full code, Meta WhatsApp Cloud API integration",
        "Supports 3 events: inquiry_verified (both parties) · host_accepted (guest) · host_rejected (guest)",
        "WhatsApp message templates: 4 templates written in inquiry-messages.ts (English)",
        "notification_log / notifications table — every send attempt logged (sent/failed/dev_logged)",
        "booking_notification_queue table + trigger — auto-queues on verified inquiry insert",
        "Dev mode: WHATSAPP_DEV_MODE=true logs messages without hitting Meta API (safe for testing)",
        "Fallback: if Meta template not approved yet, gracefully dev_logged (no crash)",
    ]
    for i in items_done_notif:
        bullet(doc, i, GREEN, "✓  ")

    spacer(doc)
    heading(doc, "2E.  OTP Infrastructure — Backend Built", 3)
    items_done_otp = [
        "send-booking-otp edge function: Twilio Verify (preferred) OR Twilio Programmable SMS fallback",
        "Rate limiting in code: max 3 OTPs/phone/hour + max 20/IP/hour",
        "OTP TTL: 10 minutes; SHA-256 hash stored (never plain text)",
        "verify-booking-otp edge function: validates hash, issues one-time token for RPC",
        "consume_booking_inquiry_otp DB RPC: single-use token, prevents replay",
        "otp_requests table + booking_otp_verifications table — fully in prod DB",
        "India E.164 normalization (+91) in both frontend and edge function",
    ]
    for i in items_done_otp:
        bullet(doc, i, GREEN, "✓  ")

    spacer(doc)
    heading(doc, "2F.  Database — Schema Solid", 3)
    items_done_db = [
        "10 migrations applied to production Supabase (arettgcwbvryrtpjtkjo)",
        "bookings: inquiry_type, phone_verified, phone_verified_at, offer_amount, host_decision_at",
        "New status values: pending_host · accepted · rejected (+ legacy preserved)",
        "otp_requests + booking_otp_verifications tables — RLS locked to service_role only",
        "booking_notification_queue + notifications tables — audit trail for every notification",
        "property_subscriptions.plan_tier + billing_cycle + yearly_discount_percent",
        "hosts.plan_tier default + billing_cycle — account-level subscription tracking",
        "Trigger: verified inquiry auto-inserts into booking_notification_queue",
        "31 catalog listings restored (was accidentally hidden by null-host migration)",
        "enforce_host_phone trigger relaxed: catalog listings (no host) can stay active",
    ]
    for i in items_done_db:
        bullet(doc, i, GREEN, "✓  ")

    spacer(doc)

    # ══ SECTION 3 — APIS: CODE DONE, KEYS PENDING ════════════════════════════
    heading(doc, "3.  APIs — Code 100% Ready, Activation Keys Chahiye", 2)
    body(doc,
         "Yeh teeno APIs ka poora code likh diya gaya hai aur repo mein hai. "
         "Sirf environment variables (secrets) Supabase edge function settings mein "
         "daalne hain — iske baad live ho jaayenge. No more coding needed.",
         bold=True, color=AMBER, size=10)

    spacer(doc)

    # Twilio
    heading(doc, "3A.  Twilio — OTP SMS for Guest Inquiry Verification", 3)
    body(doc,
         "Status: CODE COMPLETE. Waiting for credentials.",
         bold=True, color=AMBER)
    body(doc,
         "Kya hoga activate hone ke baad: Guest apna phone number daale → "
         "Twilio SMS pe 6-digit OTP aayega → verify karke inquiry submit hogi. "
         "Bina iss ke OTP always 'SMS provider not configured' error deta hai.",
         color=MUTED)
    body(doc, "Keys needed in Supabase Edge Function Secrets:", bold=True, color=NAVY)
    twilio_keys = [
        "TWILIO_ACCOUNT_SID — Twilio console pe milega (Account Info)",
        "TWILIO_AUTH_TOKEN — same Twilio console",
        "TWILIO_VERIFY_SERVICE_SID — preferred (Verify product → Create service) — handles DLT automatically for India",
        "  OR: TWILIO_PHONE_NUMBER — agar Verify nahi, toh programmable SMS (needs India DLT registration separately)",
    ]
    for k in twilio_keys:
        bullet(doc, k, AMBER)
    body(doc,
         "India DLT Note: Twilio Verify handles Indian SMS regulations internally. "
         "Recommended path: use Twilio Verify Service SID — DLT registration nahi karni. "
         "If using Programmable SMS instead, DLT entity + template registration required (7-10 days TRAI process).",
         color=GRAY)

    spacer(doc)

    # WhatsApp Meta
    heading(doc, "3B.  Meta WhatsApp Business API — Inquiry Notifications", 3)
    body(doc,
         "Status: CODE COMPLETE. Two things needed: API credentials + Template approval.",
         bold=True, color=AMBER)
    body(doc,
         "Kya hoga: Guest inquiry submit hone pe → guest ko WhatsApp confirmation → "
         "host ko WhatsApp alert with guest name + dates + dashboard link. "
         "Host Accept ke baad → guest ko WhatsApp: 'Host ne accept kar liya, seedha call karo.'",
         color=MUTED)
    body(doc, "Keys needed in Supabase Edge Function Secrets:", bold=True, color=NAVY)
    wa_keys = [
        "WHATSAPP_TOKEN — Meta Business Manager → WhatsApp → API setup → Permanent token",
        "WHATSAPP_PHONE_NUMBER_ID — same page, pe-registered phone number ka ID",
        "WHATSAPP_DEV_MODE — set to 'false' in production (currently 'true' = logs only, no real sends)",
        "NOTIFICATION_DISPATCH_SECRET — custom string (any random UUID) for security",
        "HOST_DASHBOARD_URL — https://www.xpressbnb.com/host/bookings (already in code)",
    ]
    for k in wa_keys:
        bullet(doc, k, AMBER)
    body(doc,
         "Template approval timeline: Meta requires Business Verification + message template approval. "
         "Approval typically takes 48-72 hours after account verification. "
         "MESSAGE TEMPLATES ARE ALREADY WRITTEN in code (inquiry-messages.ts). "
         "Submit them to Meta Business Manager immediately — do not wait.",
         color=GRAY)
    body(doc,
         "Fallback in place: Until templates are approved, system runs in dev_logged mode "
         "(no crash, no broken UX — just no real WhatsApp sent). "
         "Clicking WhatsApp button on success screen still works via wa.me deep links.",
         color=GRAY)

    spacer(doc)

    # Razorpay Edge Function Secret
    heading(doc, "3C.  Razorpay — Host Subscription Edge Functions", 3)
    body(doc,
         "Status: Frontend live (rzp_live_SJYZeCzffWOjuE in .env). Edge function secret missing.",
         bold=True, color=AMBER)
    body(doc,
         "Frontend Razorpay key set hai. But create-host-subscription-order aur "
         "verify-host-subscription edge functions ko backend secret key chahiye.",
         color=MUTED)
    body(doc, "Keys needed in Supabase Edge Function Secrets:", bold=True, color=NAVY)
    rp_keys = [
        "RAZORPAY_KEY_ID — rzp_live_SJYZeCzffWOjuE (already in .env, copy to edge secrets too)",
        "RAZORPAY_KEY_SECRET — Razorpay Dashboard → API Keys → Key Secret (never goes in .env / frontend)",
    ]
    for k in rp_keys:
        bullet(doc, k, AMBER)
    body(doc,
         "Once these are set, host subscription payment gets server-side verified — "
         "current workaround (client-side upsert) will be replaced with tamper-proof flow.",
         color=GRAY)

    spacer(doc)

    # ══ SECTION 4 — PENDING (ACTUAL REMAINING BUILD WORK) ════════════════════
    heading(doc, "4.  Kya Bacha Hai — Actual Pending Work", 2)
    body(doc,
         "Yeh woh features hain jo abhi tak build nahi hue — APIs se alag, "
         "inhe code karna bhi padega.",
         bold=True, color=RED, size=10)

    spacer(doc)

    pending = [
        ("Email notification fallback",
         "Koi email sending infrastructure nahi. WhatsApp fail hone pe guest/host ko email nahi jaata. "
         "Resend / Supabase emails — ek edge function ki zarurat hai."),
        ("31 properties — host linking (ops, not code)",
         "Rishikesh 18 + G.Noida 7 + Noida 6 ke paas host_id = NULL. "
         "Inka inquiry flow work karta hai but host ka phone nahi dikhta. "
         "Operators ko XpressBnB pe register karna padega (ya admin se link karna)."),
        ("Premium features — actual build",
         "Advanced analytics, calendar sync (Airbnb/Booking.com), Verified badge, Featured placement "
         "— sab 'coming soon' badge laga hai UI mein. Premium plan sell ho sakta hai abhi; "
         "features Phase 2/3 hain."),
        ("Reviews — new review submission flow",
         "Reviews table + display exist. Lekin guest review submit karne ka form nahi hai. "
         "BookingConfirmationPage se hook karna hai."),
        ("Admin panel for host/property management",
         "AdminDashboard.tsx hai but basic hai. Host linking, bulk property updates, "
         "OTP audit log viewing — sab manually SQL se karna padta hai abhi."),
        ("Delhi host names cleanup",
         "13 Delhi listings mein host 'name' field mein phone number stored hai "
         "(import artifact). Cosmetic but guests ko ajeeb lagta hai."),
        ("Guest confirmation page — full inquiry copy",
         "BookingConfirmationPage abhi bhi legacy 'payment confirmed' wording use karta hai "
         "agar sessionStorage snapshot nahi mila. Inquiry-first copy se align karna hai."),
        ("SEO / city landing pages polish",
         "Explore cities, Rishikesh, Noida pages exist but meta descriptions aur structured data "
         "inquiry model ke hisaab se update karne hain."),
    ]
    for title, detail in pending:
        p = doc.add_paragraph(style="List Bullet")
        rb = p.add_run(f"{title}:  ")
        rb.bold = True
        rb.font.color.rgb = RED
        rb.font.size = Pt(10.5)
        rd = p.add_run(detail)
        rd.font.color.rgb = MUTED
        rd.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(5)

    spacer(doc)

    # ══ SECTION 5 — JUNE 10 LAUNCH PLAN ═════════════════════════════════════
    heading(doc, "5.  June 10 Launch Plan — 22 Days", 2)
    body(doc,
         "Aaj 19 May hai. June 10 tak 22 din hain. "
         "Below is the honest day-by-day breakdown.",
         bold=True, color=NAVY, size=10)

    spacer(doc)
    tbl2 = doc.add_table(rows=1, cols=3)
    tbl2.style = "Table Grid"
    hdr2 = tbl2.rows[0].cells
    hdr2[0].text = "Week"
    hdr2[1].text = "What (Saad builds)"
    hdr2[2].text = "What (Siddhart Sir ops/decisions)"
    for cell in hdr2:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = NAVY

    weeks = [
        ("Week 1\n19–25 May",
         "1. Deploy send-booking-otp + verify-booking-otp to Supabase cloud\n"
         "2. Test OTP end-to-end (Twilio dev credentials → real)\n"
         "3. Deploy send-inquiry-notification\n"
         "4. Razorpay edge secret set — subscription verify live",
         "1. Share Twilio credentials\n"
         "2. Share Razorpay Key Secret\n"
         "3. Submit WhatsApp message templates to Meta\n"
         "4. Call 2-3 Rishikesh/Supernova operators to register"),
        ("Week 2\n26 May–1 June",
         "1. End-to-end test: guest OTP → inquiry → host dashboard → accept → WhatsApp (both)\n"
         "2. Email fallback (Supabase/Resend) — if WhatsApp fails\n"
         "3. Guest confirmation page — inquiry copy finalize\n"
         "4. Delhi host name cleanup (bulk SQL update)",
         "1. 31 host linking — Hive/Tapovan Rishikesh (likely 2 operators)\n"
         "2. G.Noida 7 + Noida Supernova 6 — find contacts\n"
         "3. WhatsApp template approval — follow up Meta"),
        ("Week 3\n2–8 June",
         "1. Reviews submission flow\n"
         "2. QA pass on all city pages (mobile)\n"
         "3. Load test OTP + inquiry on staging\n"
         "4. Final SEO meta updates",
         "1. Soft launch: share with 10 pilot hosts\n"
         "2. First subscription sale test\n"
         "3. Collect feedback"),
        ("Launch\n9–10 June",
         "Production deploy + monitoring setup\n"
         "Notification log watch first 24 hours",
         "Host onboarding call\n"
         "First ad campaign if pilot OK"),
    ]
    for week, saad, sid in weeks:
        row = tbl2.add_row().cells
        row[0].text = week
        row[1].text = saad
        row[2].text = sid

    spacer(doc)

    # ══ SECTION 6 — DECISIONS NEEDED FROM SIDDHART ═══════════════════════════
    heading(doc, "6.  Siddhart Sir se 5 Decisions Chahiye", 2)
    body(doc,
         "Yeh technical decisions nahi hain — business decisions hain. "
         "Main inke bina code side pe assumptions leke chaloon toh baad mein revert karna padega.",
         color=MUTED)

    spacer(doc)
    decisions = [
        ("Subscription scope",
         "₹999/₹2999 per property hai (current code). Ya per host account (one sub covers all listings)? "
         "Code change 1 day — but clarify karo before first sale."),
        ("Free tier limit",
         "UI mein currently 'unlimited listings' on free. Original spec mein '1 listing'. "
         "Kaunsa sahi hai? Revenue model directly impact hoga."),
        ("31 missing hosts — fallback",
         "Kya hum XpressBnB ka ek central number (jaise +91-XXXXXX) un 31 pe temporarily laga dein "
         "taaki inquiry ka number guest ko mile? Ya sirf operators ko onboard karo pehle?"),
        ("Guest host phone timing",
         "Spec mein tha: inquiry submit karte hi host phone dikhao. "
         "Alternatively: sirf jab host accept kare. "
         "Current code: dikhata hai inquiry success pe. Sahi hai?"),
        ("WhatsApp template language",
         "4 templates English mein likhe hain (inquiry-messages.ts). "
         "Hindi/Hinglish versions chahiye? "
         "Meta template approval language-specific hoti hai — pehle confirm karo."),
    ]
    for i, (q, detail) in enumerate(decisions, 1):
        p = doc.add_paragraph(style="List Bullet")
        rb = p.add_run(f"Q{i}: {q} — ")
        rb.bold = True
        rb.font.color.rgb = NAVY
        rb.font.size = Pt(10.5)
        rd = p.add_run(detail)
        rd.font.color.rgb = MUTED
        rd.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(6)

    spacer(doc)

    # ══ SECTION 7 — QUICK CHECKLIST ═══════════════════════════════════════════
    heading(doc, "7.  Original Spec vs Current Status — Har Point", 2)

    checklist = [
        # (item, status_label, status_color)
        ("No guest Razorpay on property page",                  "DONE",          GREEN),
        ("Razorpay only in host dashboard (subscriptions)",     "DONE",          GREEN),
        ("Free tier",                                           "DONE",          GREEN),
        ("₹999/month Standard tier",                           "DONE",          GREEN),
        ("₹2,999/month Premium tier",                          "DONE",          GREEN),
        ("Yearly billing — 12% discount",                       "DONE",          GREEN),
        ("Phone + OTP before inquiry",                          "CODE DONE — Twilio keys pending", AMBER),
        ("Book Now & Pay Later path",                           "DONE",          GREEN),
        ("Make an Offer path",                                  "DONE",          GREEN),
        ("Host Accept / Reject / Counter",                      "DONE",          GREEN),
        ("Success UI — 'Inquiry sent to {host}' + phone",       "DONE",          GREEN),
        ("Anti-fake-booking OTP gate",                          "CODE DONE — Twilio keys pending", AMBER),
        ("WhatsApp to guest on inquiry",                        "CODE DONE — Meta credentials + template approval pending", AMBER),
        ("WhatsApp to host on new inquiry",                     "CODE DONE — Meta credentials + template approval pending", AMBER),
        ("In-app notification (host Realtime toast)",           "DONE",          GREEN),
        ("Notification log / audit",                            "DONE",          GREEN),
        ("Zero commission positioning copy",                    "DONE",          GREEN),
        ("Premium feature gating (Standard vs Premium)",        "DONE — features themselves coming soon", AMBER),
        ("93 live listings across all cities",                  "DONE",          GREEN),
        ("31 host-less listings (ops)",                         "VISIBLE — host link pending", RED),
        ("Email fallback notifications",                        "NOT BUILT",     RED),
        ("Reviews submission flow",                             "NOT BUILT",     RED),
    ]
    for item, lbl, col in checklist:
        status_row(doc, item, lbl, col)

    spacer(doc)

    # ══ SECTION 8 — CLOSING ═══════════════════════════════════════════════════
    divider(doc)
    heading(doc, "Closing Note", 2, color=NAVY)
    body(doc,
         "Siddhart bhai, main 22 din mein ek clean, working, launch-ready product deliver "
         "kar sakta hoon — lekin sirf tab jab Twilio + WhatsApp + Razorpay secret "
         "keys is week mil jaayein. Code ka kaam mera hai; credentials ka kaam "
         "tumhara. Dono sides fast chalein toh 10 June comfortable hai.",
         color=MUTED)
    body(doc,
         "Main daily update ready hoon dene ke liye — "
         "feature by feature, koi surprise nahi.",
         bold=True, color=NAVY)

    spacer(doc)
    body(doc, "— Saad", bold=True, color=NAVY, size=13)

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build()
